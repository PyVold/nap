"""
Document Parser — Extract text from uploaded files for RAG knowledge base.
Supports: PDF, DOCX, TXT, MD, CSV, JSON, YAML
"""

import io
from typing import Tuple, Dict
from shared.logger import setup_logger

logger = setup_logger(__name__)

# Max file size: 20MB
MAX_FILE_SIZE = 20 * 1024 * 1024

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".txt", ".md", ".csv", ".json", ".yaml", ".yml",
    ".conf", ".cfg", ".log", ".xml",
}


async def extract_text_from_upload(content: bytes, filename: str) -> Tuple[str, Dict]:
    """Extract text from uploaded file bytes.

    Returns (extracted_text, metadata_dict)
    """
    if len(content) > MAX_FILE_SIZE:
        raise ValueError(f"File too large ({len(content)} bytes). Max: {MAX_FILE_SIZE} bytes.")

    ext = _get_extension(filename)
    metadata = {
        "filename": filename,
        "extension": ext,
        "size_bytes": len(content),
    }

    if ext == ".pdf":
        text, pages = _extract_pdf(content)
        metadata["pages"] = pages
    elif ext in (".docx", ".doc"):
        text, paragraphs = _extract_docx(content)
        metadata["paragraphs"] = paragraphs
    elif ext in (".txt", ".md", ".conf", ".cfg", ".log", ".csv", ".xml",
                  ".json", ".yaml", ".yml"):
        text = _extract_text(content)
    else:
        raise ValueError(
            f"Unsupported file type: {ext}. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    metadata["text_length"] = len(text)
    logger.info(f"Extracted {len(text)} chars from {filename} ({ext})")
    return text, metadata


def _get_extension(filename: str) -> str:
    """Get lowercase file extension."""
    if "." in filename:
        return "." + filename.rsplit(".", 1)[-1].lower()
    return ""


def _extract_pdf(content: bytes) -> Tuple[str, int]:
    """Extract text from PDF using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ValueError(
            "PDF support not available. Install PyPDF2: pip install PyPDF2"
        )

    reader = PdfReader(io.BytesIO(content))
    pages = len(reader.pages)
    text_parts = []

    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            text_parts.append(f"--- Page {i + 1} ---\n{page_text}")

    return "\n\n".join(text_parts), pages


def _extract_docx(content: bytes) -> Tuple[str, int]:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ValueError(
            "DOCX support not available. Install python-docx: pip install python-docx"
        )

    doc = Document(io.BytesIO(content))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading structure
            if para.style and para.style.name and "Heading" in para.style.name:
                level = para.style.name.replace("Heading ", "").strip()
                prefix = "#" * int(level) if level.isdigit() else "##"
                paragraphs.append(f"{prefix} {text}")
            else:
                paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            paragraphs.append("\n".join(rows))

    return "\n\n".join(paragraphs), len(paragraphs)


def _extract_text(content: bytes) -> str:
    """Extract text from plain text files."""
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")
